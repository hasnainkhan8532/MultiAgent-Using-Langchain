"""
Enhanced MultiToolAPI - Client-Centric System with RAG, Maps, and Chatbot
"""

from fastapi import FastAPI, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import requests
from bs4 import BeautifulSoup
# Removed legacy SDK: import google.generativeai as genai
import os
from dotenv import load_dotenv
import json
import uuid
from datetime import datetime
import sqlite3
import hashlib
# New imports for file uploads and parsing
from fastapi import UploadFile, File, Form
from io import BytesIO
# New imports for crawling
from urllib.parse import urljoin, urlparse
import time

try:
    import PyPDF2
except Exception:  # pragma: no cover
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

# Load environment variables
load_dotenv()

app = FastAPI(
    title="MultiToolAPI - Enhanced Client System",
    description="Intelligent Client Analysis with RAG, Maps, and Chatbot",
    version="2.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure Gemini (REST per https://ai.google.dev/api)
gemini_api_key = os.getenv("GEMINI_API_KEY")
gemini_model_name = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
GEMINI_GENERATE_URL = f"https://generativelanguage.googleapis.com/v1beta/models/{gemini_model_name}:generateContent"


def gemini_generate_text(prompt: str) -> str:
    if not gemini_api_key:
        raise RuntimeError("GEMINI_API_KEY not set")
    headers = {
        "x-goog-api-key": gemini_api_key,
        "Content-Type": "application/json",
    }
    body = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": prompt}],
            }
        ]
    }
    resp = requests.post(GEMINI_GENERATE_URL, headers=headers, data=json.dumps(body), timeout=30)
    data = resp.json()
    if resp.status_code >= 400:
        # Surface provider error
        raise RuntimeError(data.get("error", {}).get("message", str(data)))
    # Extract first candidate text
    text_parts: List[str] = []
    for cand in (data.get("candidates") or []):
        content = cand.get("content") or {}
        for part in content.get("parts") or []:
            if "text" in part:
                text_parts.append(part["text"])
    return "\n".join(text_parts).strip()

# Database setup
def init_database():
    conn = sqlite3.connect('multitoolapi_enhanced.db')
    cursor = conn.cursor()
    
    # Clients table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS clients (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            email TEXT NOT NULL,
            company TEXT,
            website TEXT,
            industry TEXT,
            description TEXT,
            location TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Scraped data table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS scraped_data (
            id TEXT PRIMARY KEY,
            client_id TEXT NOT NULL,
            url TEXT NOT NULL,
            title TEXT,
            meta_description TEXT,
            headings TEXT,
            links TEXT,
            text_content TEXT,
            images TEXT,
            scraped_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (client_id) REFERENCES clients (id)
        )
    ''')
    
    # RAG documents table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS rag_documents (
            id TEXT PRIMARY KEY,
            client_id TEXT NOT NULL,
            title TEXT NOT NULL,
            content TEXT NOT NULL,
            document_type TEXT,
            source TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (client_id) REFERENCES clients (id)
        )
    ''')
    
    # Maps data table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS maps_data (
            id TEXT PRIMARY KEY,
            client_id TEXT NOT NULL,
            place_id TEXT,
            name TEXT,
            address TEXT,
            phone TEXT,
            website TEXT,
            rating REAL,
            reviews_count INTEGER,
            location_lat REAL,
            location_lng REAL,
            business_hours TEXT,
            photos TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (client_id) REFERENCES clients (id)
        )
    ''')
    
    # Chat conversations table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS chat_conversations (
            id TEXT PRIMARY KEY,
            client_id TEXT NOT NULL,
            user_message TEXT NOT NULL,
            bot_response TEXT NOT NULL,
            context_data TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (client_id) REFERENCES clients (id)
        )
    ''')
    
    conn.commit()
    conn.close()

# Initialize database
init_database()

# Pydantic models
class Client(BaseModel):
    name: str
    email: str
    company: Optional[str] = None
    website: Optional[str] = None
    industry: Optional[str] = None
    description: Optional[str] = None
    location: Optional[str] = None

class ClientResponse(BaseModel):
    id: str
    name: str
    email: str
    company: Optional[str] = None
    website: Optional[str] = None
    industry: Optional[str] = None
    description: Optional[str] = None
    location: Optional[str] = None
    created_at: str
    updated_at: str

class ScrapingRequest(BaseModel):
    client_id: str
    url: str
    strategy: str = "requests"

class CrawlRequest(BaseModel):
    client_id: str
    start_url: str
    max_pages: int = 10
    max_depth: int = 2
    same_domain_only: bool = True
    delay: float = 0.5

class RAGDocument(BaseModel):
    client_id: str
    title: str
    content: str
    document_type: str = "general"
    source: Optional[str] = None

class MapsRequest(BaseModel):
    client_id: str
    query: str
    location: Optional[str] = None

class ChatMessage(BaseModel):
    client_id: str
    message: str

class AnalysisRequest(BaseModel):
    client_id: str
    requirements: str
    goals: List[str]

# Database helper functions
def get_db_connection():
    return sqlite3.connect('multitoolapi_enhanced.db')

def get_client_data(client_id: str) -> Dict[str, Any]:
    """Get all data for a specific client"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get client info
    cursor.execute('SELECT * FROM clients WHERE id = ?', (client_id,))
    client = cursor.fetchone()
    if not client:
        conn.close()
        return None
    
    # Get scraped data
    cursor.execute('SELECT * FROM scraped_data WHERE client_id = ? ORDER BY scraped_at DESC', (client_id,))
    scraped_data = cursor.fetchall()
    
    # Get RAG documents
    cursor.execute('SELECT * FROM rag_documents WHERE client_id = ? ORDER BY created_at DESC', (client_id,))
    rag_documents = cursor.fetchall()
    
    # Get maps data
    cursor.execute('SELECT * FROM maps_data WHERE client_id = ? ORDER BY created_at DESC', (client_id,))
    maps_data = cursor.fetchall()
    
    # Get chat history
    cursor.execute('SELECT * FROM chat_conversations WHERE client_id = ? ORDER BY created_at DESC LIMIT 10', (client_id,))
    chat_history = cursor.fetchall()
    
    conn.close()
    
    return {
        'client': client,
        'scraped_data': scraped_data,
        'rag_documents': rag_documents,
        'maps_data': maps_data,
        'chat_history': chat_history
    }

# API Endpoints
@app.get("/")
async def root():
    return {
        "message": "MultiToolAPI - Enhanced Client System",
        "version": "2.0.0",
        "status": "running"
    }

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "gemini_configured": bool(gemini_api_key),
        "gemini_model": os.getenv("GEMINI_MODEL", "gemini-2.5-flash"),
        "database": "sqlite"
    }

# New: AI status check
@app.get("/ai/status")
async def ai_status():
    if not gemini_api_key:
        return {"ok": False, "configured": False, "provider": "gemini", "message": "GEMINI_API_KEY not set"}
    # Try a very small request via REST
    try:
        _ = gemini_generate_text("ping")
        return {"ok": True, "configured": True, "provider": "gemini", "model": os.getenv("GEMINI_MODEL", "gemini-2.5-flash")}
    except Exception as e:
        return {"ok": False, "configured": True, "provider": "gemini", "model": os.getenv("GEMINI_MODEL", "gemini-2.5-flash"), "error": str(e)}

# Client Management
@app.post("/clients/", response_model=ClientResponse)
async def create_client(client: Client):
    """Create a new client"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    client_id = str(uuid.uuid4())
    now = datetime.now().isoformat()
    
    cursor.execute('''
        INSERT INTO clients (id, name, email, company, website, industry, description, location, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (client_id, client.name, client.email, client.company, client.website, 
          client.industry, client.description, client.location, now, now))
    
    conn.commit()
    conn.close()
    
    return ClientResponse(
        id=client_id,
        name=client.name,
        email=client.email,
        company=client.company,
        website=client.website,
        industry=client.industry,
        description=client.description,
        location=client.location,
        created_at=now,
        updated_at=now
    )

@app.get("/clients/", response_model=List[ClientResponse])
async def get_clients():
    """Get all clients"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM clients ORDER BY created_at DESC')
    clients = cursor.fetchall()
    conn.close()
    
    return [
        ClientResponse(
            id=client[0],
            name=client[1],
            email=client[2],
            company=client[3],
            website=client[4],
            industry=client[5],
            description=client[6],
            location=client[7],
            created_at=client[8],
            updated_at=client[9]
        ) for client in clients
    ]

@app.get("/clients/{client_id}", response_model=ClientResponse)
async def get_client(client_id: str):
    """Get a specific client"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM clients WHERE id = ?', (client_id,))
    client = cursor.fetchone()
    conn.close()
    
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    return ClientResponse(
        id=client[0],
        name=client[1],
        email=client[2],
        company=client[3],
        website=client[4],
        industry=client[5],
        description=client[6],
        location=client[7],
        created_at=client[8],
        updated_at=client[9]
    )

# Web Scraping aligned with clients
@app.post("/scraping/scrape")
async def scrape_website(request: ScrapingRequest):
    """Scrape a website for a specific client"""
    try:
        response = requests.get(request.url, timeout=30)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Extract information
        scraped_data = {
            "url": request.url,
            "title": soup.title.string if soup.title else "",
            "meta_description": soup.find('meta', attrs={'name': 'description'}).get('content', '') if soup.find('meta', attrs={'name': 'description'}) else "",
            "headings": [h.get_text().strip() for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])],
            "links": [{"text": a.get_text().strip(), "href": a.get('href', '')} for a in soup.find_all('a', href=True)],
            "text_content": soup.get_text().strip(),
            "images": [{"src": img.get('src', ''), "alt": img.get('alt', '')} for img in soup.find_all('img')]
        }
        
        # Save to database
        conn = get_db_connection()
        cursor = conn.cursor()
        
        data_id = str(uuid.uuid4())
        cursor.execute('''
            INSERT INTO scraped_data (id, client_id, url, title, meta_description, headings, links, text_content, images)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (data_id, request.client_id, scraped_data['url'], scraped_data['title'],
              scraped_data['meta_description'], json.dumps(scraped_data['headings']),
              json.dumps(scraped_data['links']), scraped_data['text_content'],
              json.dumps(scraped_data['images'])))
        
        conn.commit()
        conn.close()
        
        return {
            "message": "Website scraped successfully",
            "data": scraped_data,
            "client_id": request.client_id
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scraping failed: {str(e)}")

@app.post("/scraping/crawl")
async def crawl_website(request: CrawlRequest):
    """Crawl a website starting from start_url and save up to max_pages within depth limits."""
    start = request.start_url
    parsed_start = urlparse(start)
    start_domain = parsed_start.netloc

    to_visit = [(start, 0)]
    visited = set()
    saved = 0
    errors = 0

    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        while to_visit and saved < request.max_pages:
            current_url, depth = to_visit.pop(0)
            if current_url in visited:
                continue
            visited.add(current_url)

            try:
                resp = requests.get(current_url, timeout=30)
                resp.raise_for_status()
            except Exception:
                errors += 1
                continue

            soup = BeautifulSoup(resp.content, 'html.parser')
            # Extract
            title = soup.title.string if soup.title else ""
            meta_description = soup.find('meta', attrs={'name': 'description'}).get('content', '') if soup.find('meta', attrs={'name': 'description'}) else ""
            headings = [h.get_text().strip() for h in soup.find_all(['h1','h2','h3','h4','h5','h6'])]
            links = [{"text": a.get_text().strip(), "href": urljoin(current_url, a.get('href', ''))} for a in soup.find_all('a', href=True)]
            text_content = soup.get_text().strip()
            images = [{"src": urljoin(current_url, img.get('src', '')), "alt": img.get('alt', '')} for img in soup.find_all('img')]

            # Save page
            data_id = str(uuid.uuid4())
            cursor.execute('''
                INSERT INTO scraped_data (id, client_id, url, title, meta_description, headings, links, text_content, images)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (data_id, request.client_id, current_url, title, meta_description,
                  json.dumps(headings), json.dumps(links), text_content, json.dumps(images)))
            conn.commit()
            saved += 1

            # Enqueue next links
            if depth < request.max_depth:
                for a in soup.find_all('a', href=True):
                    next_href = urljoin(current_url, a.get('href'))
                    parsed_next = urlparse(next_href)
                    if parsed_next.scheme not in ("http", "https"):
                        continue
                    if request.same_domain_only and parsed_next.netloc != start_domain:
                        continue
                    if next_href not in visited:
                        to_visit.append((next_href, depth + 1))

            if request.delay and request.delay > 0:
                time.sleep(request.delay)

        return {
            "message": "Crawl completed",
            "start_url": start,
            "pages_saved": saved,
            "errors": errors,
            "visited": len(visited)
        }
    finally:
        conn.close()

@app.get("/clients/{client_id}/scraped-data")
async def get_client_scraped_data(client_id: str):
    """Get all scraped data for a specific client"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM scraped_data WHERE client_id = ? ORDER BY scraped_at DESC', (client_id,))
    data = cursor.fetchall()
    conn.close()
    
    return {
        "client_id": client_id,
        "scraped_data": [
            {
                "id": row[0],
                "url": row[2],
                "title": row[3],
                "meta_description": row[4],
                "headings": json.loads(row[5]) if row[5] else [],
                "links": json.loads(row[6]) if row[6] else [],
                "text_content": row[7],
                "images": json.loads(row[8]) if row[8] else [],
                "scraped_at": row[9]
            } for row in data
        ]
    }

# RAG System
@app.post("/rag/add-document")
async def add_rag_document(document: RAGDocument):
    """Add a document to RAG system for a specific client"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    doc_id = str(uuid.uuid4())
    now = datetime.now().isoformat()
    
    cursor.execute('''
        INSERT INTO rag_documents (id, client_id, title, content, document_type, source, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (doc_id, document.client_id, document.title, document.content,
          document.document_type, document.source, now))
    
    conn.commit()
    conn.close()
    
    return {
        "message": "Document added to RAG system",
        "document_id": doc_id,
        "client_id": document.client_id
    }

# New: upload endpoint to accept PDF/TXT/DOCX
@app.post("/rag/upload")
async def upload_rag_document(
    client_id: str = Form(...),
    document_type: str = Form("general"),
    source: Optional[str] = Form(None),
    file: UploadFile = File(...),
):
    """Upload a file (pdf/txt/docx) and store its text as a RAG document."""
    # Read file bytes
    file_bytes = await file.read()
    filename = file.filename or f"upload-{uuid.uuid4()}"
    ext = os.path.splitext(filename)[1].lower()

    # Extract text based on extension
    text_content = ""
    try:
        if ext == ".pdf":
            if not PyPDF2:
                raise HTTPException(status_code=500, detail="PDF support not available on server")
            reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            parts: List[str] = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    continue
            text_content = "\n".join(parts).strip()
        elif ext in (".txt", ".md", ".csv", ".log"):
            try:
                text_content = file_bytes.decode("utf-8", errors="ignore")
            except Exception:
                text_content = file_bytes.decode("latin-1", errors="ignore")
        elif ext in (".docx",):
            if not DocxDocument:
                raise HTTPException(status_code=500, detail="DOCX support not available on server")
            doc = DocxDocument(BytesIO(file_bytes))
            text_content = "\n".join(p.text for p in doc.paragraphs)
        else:
            # Fallback: try utf-8 decode
            try:
                text_content = file_bytes.decode("utf-8", errors="ignore")
            except Exception:
                text_content = ""
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")

    if not text_content or not text_content.strip():
        raise HTTPException(status_code=400, detail="No text could be extracted from the file")

    # Derive title from filename
    title = os.path.splitext(os.path.basename(filename))[0]

    # Persist into database
    conn = get_db_connection()
    cursor = conn.cursor()
    doc_id = str(uuid.uuid4())
    now = datetime.now().isoformat()
    cursor.execute('''
        INSERT INTO rag_documents (id, client_id, title, content, document_type, source, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (doc_id, client_id, title, text_content, document_type, source or filename, now))
    conn.commit()
    conn.close()

    return {
        "message": "File uploaded and document added",
        "document_id": doc_id,
        "client_id": client_id,
        "filename": filename,
        "chars": len(text_content)
    }

@app.get("/clients/{client_id}/rag-documents")
async def get_client_rag_documents(client_id: str):
    """Get all RAG documents for a specific client"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM rag_documents WHERE client_id = ? ORDER BY created_at DESC', (client_id,))
    documents = cursor.fetchall()
    conn.close()
    
    return {
        "client_id": client_id,
        "documents": [
            {
                "id": doc[0],
                "title": doc[2],
                "content": doc[3],
                "document_type": doc[4],
                "source": doc[5],
                "created_at": doc[6]
            } for doc in documents
        ]
    }

# Google Maps Integration
@app.post("/maps/search")
async def search_places(request: MapsRequest):
    """Search for places using Google Maps API"""
    google_api_key = os.getenv("GOOGLE_MAPS_API_KEY")
    if not google_api_key:
        raise HTTPException(status_code=500, detail="Google Maps API key not configured")
    
    try:
        # Google Places API search
        url = "https://maps.googleapis.com/maps/api/place/textsearch/json"
        params = {
            "query": request.query,
            "key": google_api_key
        }
        
        if request.location:
            params["location"] = request.location
            params["radius"] = "50000"  # 50km radius
        
        response = requests.get(url, params=params)
        data = response.json()
        
        if data.get("status") != "OK":
            raise HTTPException(status_code=400, detail=f"Places API error: {data.get('error_message', 'Unknown error')}")
        
        # Save results to database
        conn = get_db_connection()
        cursor = conn.cursor()
        
        saved_places = []
        for place in data.get("results", [])[:5]:  # Limit to 5 results
            place_id = str(uuid.uuid4())
            
            # Extract place details
            name = place.get("name", "")
            address = place.get("formatted_address", "")
            rating = place.get("rating", 0)
            reviews_count = place.get("user_ratings_total", 0)
            location = place.get("geometry", {}).get("location", {})
            lat = location.get("lat", 0)
            lng = location.get("lng", 0)
            
            cursor.execute('''
                INSERT INTO maps_data (id, client_id, place_id, name, address, rating, reviews_count, location_lat, location_lng)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (place_id, request.client_id, place.get("place_id", ""), name, address,
                  rating, reviews_count, lat, lng))
            
            saved_places.append({
                "id": place_id,
                "name": name,
                "address": address,
                "rating": rating,
                "reviews_count": reviews_count,
                "location": {"lat": lat, "lng": lng}
            })
        
        conn.commit()
        conn.close()
        
        return {
            "message": "Places search completed",
            "client_id": request.client_id,
            "places": saved_places
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Maps search failed: {str(e)}")

@app.get("/clients/{client_id}/maps-data")
async def get_client_maps_data(client_id: str):
    """Get all maps data for a specific client"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM maps_data WHERE client_id = ? ORDER BY created_at DESC', (client_id,))
    data = cursor.fetchall()
    conn.close()
    
    return {
        "client_id": client_id,
        "places": [
            {
                "id": row[0],
                "place_id": row[2],
                "name": row[3],
                "address": row[4],
                "phone": row[5],
                "website": row[6],
                "rating": row[7],
                "reviews_count": row[8],
                "location": {"lat": row[9], "lng": row[10]},
                "business_hours": row[11],
                "photos": row[12],
                "created_at": row[13]
            } for row in data
        ]
    }

# Intelligent Chatbot
@app.post("/chat/analyze")
async def chat_analyze(message: ChatMessage):
    """Intelligent chatbot that analyzes all client data"""
    
    # Get all client data
    client_data = get_client_data(message.client_id)
    if not client_data:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Prepare context from all data sources
    context = prepare_analysis_context(client_data)
    
    # Generate response using Gemini or fallback
    if gemini_api_key:
        try:
            prompt = f"""
            You are an intelligent business analyst chatbot. Analyze the following client data and respond to the user's question.
            
            CLIENT: {client_data['client'][1]} ({client_data['client'][3]})
            INDUSTRY: {client_data['client'][5]}
            
            AVAILABLE DATA:
            {context}
            
            USER QUESTION: {message.message}
            
            Provide a comprehensive analysis based on all available data sources. Be specific and actionable.
            """
            
            ai_text = gemini_generate_text(prompt)
            bot_response = ai_text
            analysis_type = "AI-powered"
        except Exception as e:
            # Surface concise guidance if the API key is blocked or not enabled
            guidance = ""
            err = str(e)
            if "blocked" in err.lower() or "api_key" in err.lower() or "expired" in err.lower():
                guidance = "\n\nNote: The Gemini API key appears invalid/blocked. Ensure the Gemini API (formerly Generative Language) is enabled for this project and that the key allows access to `generativelanguage.googleapis.com` (or remove API restrictions). Then restart the backend."
            bot_response = generate_fallback_analysis(client_data, message.message) + guidance
            analysis_type = "Rule-based"
    else:
        bot_response = generate_fallback_analysis(client_data, message.message)
        analysis_type = "Rule-based"
    
    # Save conversation
    conn = get_db_connection()
    cursor = conn.cursor()
    
    conv_id = str(uuid.uuid4())
    cursor.execute('''
        INSERT INTO chat_conversations (id, client_id, user_message, bot_response, context_data)
        VALUES (?, ?, ?, ?, ?)
    ''', (conv_id, message.client_id, message.message, bot_response, context))
    
    conn.commit()
    conn.close()
    
    return {
        "message": "Analysis completed",
        "response": bot_response,
        "analysis_type": analysis_type,
        "client_id": message.client_id
    }

def prepare_analysis_context(client_data):
    """Prepare context from all client data sources"""
    context_parts = []
    
    # Client basic info
    client = client_data['client']
    context_parts.append(f"Client: {client[1]} ({client[3]}) - {client[5]}")
    
    # Scraped data
    if client_data['scraped_data']:
        context_parts.append("\nSCRAPED WEBSITE DATA:")
        for data in client_data['scraped_data'][:3]:  # Last 3 scraped sites
            context_parts.append(f"- {data[3]} ({data[2]})")
            if data[4]:  # meta description
                context_parts.append(f"  Description: {data[4]}")
    
    # RAG documents
    if client_data['rag_documents']:
        context_parts.append("\nKNOWLEDGE BASE:")
        for doc in client_data['rag_documents'][:3]:  # Last 3 documents
            context_parts.append(f"- {doc[2]} ({doc[4]})")
            context_parts.append(f"  Content: {doc[3][:200]}...")
    
    # Maps data
    if client_data['maps_data']:
        context_parts.append("\nLOCATION DATA:")
        for place in client_data['maps_data'][:3]:  # Last 3 places
            context_parts.append(f"- {place[3]} ({place[4]}) - Rating: {place[7]}")
    
    return "\n".join(context_parts)

def generate_fallback_analysis(client_data, user_message):
    """Generate fallback analysis when AI is not available"""
    client = client_data['client']
    
    analysis = f"""
# CLIENT ANALYSIS: {client[1]} ({client[3]})

## Current Data Overview:
- **Industry:** {client[5]}
- **Website:** {client[4] or 'Not provided'}
- **Scraped Sites:** {len(client_data['scraped_data'])}
- **Knowledge Documents:** {len(client_data['rag_documents'])}
- **Location Data:** {len(client_data['maps_data'])}

## Analysis Based on Available Data:

### Website Performance:
"""
    
    if client_data['scraped_data']:
        latest_site = client_data['scraped_data'][0]
        analysis += f"- Latest scraped site: {latest_site[3]} ({latest_site[2]})\n"
        analysis += f"- Meta description: {'Present' if latest_site[4] else 'Missing'}\n"
        analysis += f"- Headings found: {len(json.loads(latest_site[5]) if latest_site[5] else [])}\n"
    else:
        analysis += "- No website data available yet\n"
    
    analysis += f"""
### Knowledge Base:
- Documents available: {len(client_data['rag_documents'])}
- Document types: {', '.join(set([doc[4] for doc in client_data['rag_documents']])) if client_data['rag_documents'] else 'None'}

### Location Intelligence:
- Places found: {len(client_data['maps_data'])}
- Average rating: {sum([place[7] for place in client_data['maps_data']]) / len(client_data['maps_data']) if client_data['maps_data'] else 'N/A'}

## Recommendations:
1. **Data Collection:** {'Good data coverage' if len(client_data['scraped_data']) > 0 else 'Start by scraping the client website'}
2. **Knowledge Management:** {'Well documented' if len(client_data['rag_documents']) > 2 else 'Add more knowledge documents'}
3. **Location Strategy:** {'Location data available' if client_data['maps_data'] else 'Consider adding location-based insights'}

---
*This analysis was generated using available data. For AI-powered insights, please configure your Gemini API key.*
"""
    
    return analysis

# Dashboard endpoint
@app.get("/dashboard")
async def get_dashboard():
    """Get dashboard data for all clients"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get client counts
    cursor.execute('SELECT COUNT(*) FROM clients')
    total_clients = cursor.fetchone()[0]
    
    cursor.execute('SELECT COUNT(*) FROM scraped_data')
    total_scraped = cursor.fetchone()[0]
    
    cursor.execute('SELECT COUNT(*) FROM rag_documents')
    total_documents = cursor.fetchone()[0]
    
    cursor.execute('SELECT COUNT(*) FROM maps_data')
    total_places = cursor.fetchone()[0]
    
    # Get recent activity
    cursor.execute('''
        SELECT c.name, c.company, sd.title, sd.scraped_at
        FROM clients c
        LEFT JOIN scraped_data sd ON c.id = sd.client_id
        ORDER BY sd.scraped_at DESC
        LIMIT 10
    ''')
    recent_activity = cursor.fetchall()
    
    conn.close()
    
    return {
        "total_clients": total_clients,
        "total_scraped_sites": total_scraped,
        "total_documents": total_documents,
        "total_places": total_places,
        "recent_activity": [
            {
                "client": activity[0],
                "company": activity[1],
                "action": f"Scraped: {activity[2]}" if activity[2] else "No recent activity",
                "timestamp": activity[3]
            } for activity in recent_activity
        ]
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
